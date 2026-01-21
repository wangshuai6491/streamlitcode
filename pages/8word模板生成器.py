import streamlit as st
from docx import Document
import os
import re
import json
import pandas as pd
from io import BytesIO
import tempfile
import zipfile
import openpyxl
from openpyxl.utils import column_index_from_string, get_column_letter

# 解析word模板变量的函数
def analyze_template(template_path):
    """
    分析Word模板，提取所有占位符信息
    
    :param template_path: Word模板文件路径(.docx格式)
    :return: 包含普通变量和表格信息的字典
    """
    # 加载Word模板
    doc = Document(template_path)
    
    # 正则表达式匹配占位符 {xxx}
    placeholder_pattern = re.compile(r'{([^{}]*)}')
    all_placeholders = set()
    
    # 遍历所有段落，提取占位符
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            matches = placeholder_pattern.findall(run.text)
            for match in matches:
                all_placeholders.add(f"{{{match}}}")
    
    # 遍历所有表格，提取占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = placeholder_pattern.findall(cell.text)
                for match in matches:
                    all_placeholders.add(f"{{{match}}}")
    
    # 分类占位符：普通变量和表格变量
    normal_placeholders = []
    table_info = {}
    
    for ph in all_placeholders:
        ph_content = ph[1:-1]  # 去掉 {} 括号
        if ph_content.startswith("table."):
            # 分割表格占位符：table.表格名.列名
            parts = ph_content.split(".")
            if len(parts) == 3:
                table_name = parts[1]
                column_name = parts[2]
                
                # 添加到表格信息字典
                if table_name not in table_info:
                    table_info[table_name] = {
                        "columns": set(),
                        "placeholders": []
                    }
                table_info[table_name]["columns"].add(column_name)
                table_info[table_name]["placeholders"].append(ph)
        else:
            normal_placeholders.append(ph)
    
    # 转换columns集合为列表，方便使用
    for table_name in table_info:
        table_info[table_name]["columns"] = sorted(list(table_info[table_name]["columns"]))
    
    return {
        "normal_placeholders": sorted(normal_placeholders),
        "table_info": table_info
    }

# 生成word的主要函数
def generate_contract_doc(template_path, data, doc_index=1):
    """
    基于Word模板生成文档
    
    :param template_path: Word模板文件路径(.docx格式)
    :param data: 文档数据，字典格式，包含text_data和tables_data
    :param doc_index: 文档序号，用于生成文件名
    
    :return: 生成的文档对象和文件名
    """
    # 加载Word模板
    doc = Document(template_path)

    # 分析模板，获取占位符信息
    template_analysis = analyze_template(template_path)
    
    # ========== 核心功能1：替换所有【普通文本占位符】 {xxx} ==========
    # 遍历文档中所有段落，替换文本占位符
    for paragraph in doc.paragraphs:
        # 遍历段落中的所有文本块，避免替换格式丢失
        for run in paragraph.runs:
            for placeholder, value in data["text_data"].items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

    # ========== 核心功能2：动态填充【表格】并插入多行 ==========
    # 遍历所有表格，处理包含table.xxx.xxx格式占位符的表格
    for table in doc.tables:
        # 检查表格是否包含模板行（包含table.xxx.xxx占位符）
        template_row = None
        table_name_in_row = None
        
        for row_idx, row in enumerate(table.rows):
            row_text = " ".join([cell.text for cell in row.cells])
            if "{table." in row_text:
                # 提取表格名称
                matches = re.findall(r'{table\.([^\.}]+)\.', row_text)
                if matches:
                    table_name_in_row = matches[0]
                    template_row = row
                    break
        
        if template_row and table_name_in_row:
            # 检查是否有对应的数据
            if table_name_in_row not in data.get("tables_data", {}):
                st.warning(f"表格 {table_name_in_row} 没有对应的数据，跳过该表格")
                continue
            
            # 获取表格数据
            table_data = data["tables_data"][table_name_in_row]
            
            # 获取模板行中各占位符的位置映射
            placeholder_positions = {}
            template_cells = template_row.cells
            for i, cell in enumerate(template_cells):
                cell_text = cell.text.strip()
                if cell_text.startswith("{table.") and cell_text.endswith("}"):
                    placeholder_positions[cell_text] = i
            
            # 遍历表格数据，逐行插入
            for index, row_data in enumerate(table_data):
                if index == 0:
                    # 第一行数据：直接使用模板行
                    row = template_row.cells
                    # 保存第一行的单元格作为格式模板
                    format_template_cells = template_row.cells
                else:
                    # 第2-N行数据：在模板行下方新增一行
                    new_row = table.add_row()
                    row = new_row.cells
                    
                    # 复制第一行的格式到新行
                    # 复制行属性
                    new_row.height = template_row.height
                    
                    # 复制每个单元格的格式
                    for i, (new_cell, template_cell) in enumerate(zip(row, format_template_cells)):
                        # 复制单元格宽度
                        new_cell.width = template_cell.width
                        
                        # 复制段落格式
                        for new_paragraph, template_paragraph in zip(new_cell.paragraphs, template_cell.paragraphs):
                            # 复制段落对齐方式
                            new_paragraph.alignment = template_paragraph.alignment
                            
                            # 复制段落样式
                            if template_paragraph.style:
                                new_paragraph.style = template_paragraph.style
                            
                            # 复制段落中的运行元素格式
                            for new_run, template_run in zip(new_paragraph.runs, template_paragraph.runs):
                                # 复制字体格式
                                new_run.font.name = template_run.font.name
                                new_run.font.size = template_run.font.size
                                new_run.font.bold = template_run.font.bold
                                new_run.font.italic = template_run.font.italic
                                new_run.font.underline = template_run.font.underline
                                if template_run.font.color.rgb:
                                    new_run.font.color.rgb = template_run.font.color.rgb
                                
                                # 复制其他字体属性
                                new_run.font.superscript = template_run.font.superscript
                                new_run.font.subscript = template_run.font.subscript
                                new_run.font.shadow = template_run.font.shadow
                                new_run.font.strike = template_run.font.strike
                                new_run.font.double_strike = template_run.font.double_strike
                
                # 给当前行的每个单元格赋值
                for placeholder, pos in placeholder_positions.items():
                    ph_content = placeholder[1:-1]  # 去掉 {} 括号
                    # 分割占位符：table.表格名.列名
                    parts = ph_content.split(".")
                    if len(parts) == 3:
                        column_name = parts[2]
                        
                        if column_name == "序号":
                            row[pos].text = str(index + 1)  # 序号从1开始
                        elif column_name in row_data:
                            row[pos].text = str(row_data[column_name])
                        else:
                            # 如果数据中没有该列，留空或使用默认值
                            row[pos].text = ""

    # 提取提供的数据中有没有文件名
    if "{文件名}" in data.get("text_data", {}) and data["text_data"]["{文件名}"] is not None and data["text_data"]["{文件名}"].strip():
        file_name = data["text_data"]["{文件名}"]
        if not file_name.endswith(".docx"):
            file_name += ".docx"
    # 没有则按顺序号命名
    else:
        file_name = f"{doc_index}.docx"
    
    # 保存到BytesIO对象
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    return doc_stream, file_name

# 侧边栏下载模板的函数
def cebianlan():
    # 模板下载
    st.sidebar.subheader("模板下载")
    # 使用相对路径，相对于应用运行目录
    template_file_path = os.path.join("static", "word批量生成模板.zip")
    if os.path.exists(template_file_path):
        with open(template_file_path, "rb") as f:
            st.sidebar.download_button(
                label="下载示例模板",
                data=f,
                file_name="word批量生成模板.zip",
                mime="application/zip"
            )
    else:
        st.sidebar.warning(f"示例模板文件不存在，路径：{template_file_path}")

# 展示模板占位符信息
def display_template_info(uploaded_template):
    """
    显示Word模板中的占位符信息
    
    :param uploaded_template: 上传的Word模板文件对象
    :param tmp_path: 临时文件路径
    """
    try:
        template_info = analyze_template(uploaded_template)
        
        with st.expander("📋 word占位符解析", expanded=False):
            # 显示普通变量
            st.subheader("📝 普通变量")
            if template_info["normal_placeholders"]:
                # 将所有普通变量拼接成一个字符串，每个占一行
                normal_vars_text = "\n".join(template_info["normal_placeholders"])
                st.code(normal_vars_text, language="plaintext")
            else:
                st.info("无普通变量")
            
            # 显示表格变量
            st.subheader("📊 表格变量")
            if template_info["table_info"]:
                # 构建表格变量文本，按表格分组，每个占一行
                table_vars_lines = []
                for table_name, info in template_info["table_info"].items():
                    table_vars_lines.append(f"=== {table_name} ===")
                    table_vars_lines.extend(info["placeholders"])
                    table_vars_lines.append("")  # 表格之间空一行
                # 移除最后一个空行
                if table_vars_lines and table_vars_lines[-1] == "":
                    table_vars_lines.pop()
                # 拼接成一个字符串
                table_vars_text = "\n".join(table_vars_lines)
                st.code(table_vars_text, language="plaintext")
            else:
                st.info("无表格变量")
    except Exception as e:
        st.error(f"模板分析失败：{e}")


# 把excel解析为特定格式json
def excel_to_json(excel_file, text_data_start, tables_data_mapping, sheetname=None):
    """
    将Excel文件转换为指定格式的JSON格式
    
    参数:
    excel_file: str - Excel文件路径
    text_data_start: str - text_data的左上角单元格，如"A1"
    tables_data_mapping: dict - tables_data的表名及其对应的左上角单元格，如{"家庭成员": "A10"}
    sheetname: str - 可选，指定要处理的sheet名称，如果不提供则处理所有sheet
    """
    # 读取Excel文件
    workbook = openpyxl.load_workbook(excel_file)
    
    result = []
    
    # 确定要处理的sheet列表
    sheet_names = [sheetname] if sheetname and sheetname in workbook.sheetnames else workbook.sheetnames
    
    # 遍历要处理的sheet
    for sheet_name in sheet_names:
        sheet = workbook[sheet_name]
        
        # 解析text_data起始单元格
        text_col = column_index_from_string(text_data_start[0])
        text_row = int(text_data_start[1:])
        
        if tables_data_mapping:
            # 有表格数据的情况：第一行为键，第二行为值，每个sheet生成一个文档
            # 处理text_data
            text_data = {}
            
            # text_data的键行和值行
            key_row = text_row
            value_row = text_row + 1
            current_col = text_col
            
            # 遍历所有列，直到遇到空单元格
            while True:
                key_cell = sheet.cell(row=key_row, column=current_col)
                value_cell = sheet.cell(row=value_row, column=current_col)
                
                # 检查是否为空单元格
                if key_cell.value is None:
                    break
                
                # 添加到text_data，键用{}包起来
                text_data[f"{{{key_cell.value}}}"] = value_cell.value
                current_col += 1
            
            # 处理tables_data
            tables_data = {}
            
            for table_name, table_start in tables_data_mapping.items():
                # 解析表格起始单元格
                table_col = column_index_from_string(table_start[0])
                table_row = int(table_start[1:])
                
                # 读取表头（起始行）
                headers = []
                header_row = table_row
                current_col = table_col
                
                while True:
                    header_cell = sheet.cell(row=header_row, column=current_col)
                    if header_cell.value is None:
                        break
                    headers.append(header_cell.value)
                    current_col += 1
                
                # 读取数据行（从起始行+1开始）
                table_data = []
                data_start_row = table_row + 1
                current_row = data_start_row
                
                while True:
                    # 检查第一列是否为空，为空则结束
                    first_cell = sheet.cell(row=current_row, column=table_col)
                    if first_cell.value is None:
                        break
                    
                    # 读取当前行数据
                    row_data = {}
                    for i, header in enumerate(headers):
                        value_cell = sheet.cell(row=current_row, column=table_col + i)
                        row_data[header] = value_cell.value
                    
                    table_data.append(row_data)
                    current_row += 1
                
                # 添加到tables_data
                tables_data[table_name] = table_data
            
            # 构建当前sheet的JSON对象
            sheet_data = {
                "text_data": text_data,
                "tables_data": tables_data
            }
            
            result.append(sheet_data)
        else:
            # 只有普通变量的情况：第一行永远为键，从第二行开始为值，如果后面还有第3、4、5....行也是值
            # 读取第一行的所有键
            keys = []
            current_col = text_col
            
            # 遍历第一行，获取所有键
            while True:
                key_cell = sheet.cell(row=text_row, column=current_col)
                if key_cell.value is None or pd.isna(key_cell.value):
                    break
                keys.append(str(key_cell.value).strip())
                current_col += 1
            
            # 如果没有键，直接返回
            if not keys:
                return result
            
            # 从第二行开始读取值，每行生成一个文档
            current_row = text_row + 1
            
            while True:
                # 检查第一列是否为空，为空则结束
                first_cell = sheet.cell(row=current_row, column=text_col)
                if first_cell.value is None or pd.isna(first_cell.value):
                    break
                
                # 创建当前行的text_data
                text_data = {}
                
                # 遍历所有键，读取对应列的值
                for i, key in enumerate(keys):
                    value_cell = sheet.cell(row=current_row, column=text_col + i)
                    value = value_cell.value
                    
                    # 只有当值不为空时才添加到text_data
                    if not (pd.isna(value) or (isinstance(value, str) and str(value).strip() == "")):
                        # 添加到text_data，键用{}包起来
                        text_data[f"{{{key}}}"] = str(value).strip()
                
                # 只有当text_data不为空时才添加到结果
                if text_data:
                    # 构建当前行的JSON对象
                    row_data = {
                        "text_data": text_data,
                        "tables_data": {}
                    }
                    
                    result.append(row_data)
                
                current_row += 1
    
    return result

# 只有普通变量的数据解析函数，最终处理为json
def only_normal_vars():
    """
    处理只有普通变量的情况：用户上传Excel后，选择表格，以A1单元格为基础读取连续区域的内容并解析为JSON格式
    """
    # 上传Excel文件
    uploaded_excel = st.file_uploader("上传Excel文件（仅含普通变量）", type=["xlsx", "xls"], key="only_normal_vars_excel")
    
    if uploaded_excel:
        try:
            # 读取Excel文件获取sheet列表
            excel_data = pd.ExcelFile(uploaded_excel)
            
            # 让用户选择表格（工作表）
            selected_sheet = st.selectbox(
                "选择工作表",
                excel_data.sheet_names,
                key="only_normal_vars_sheet"
            )
        
            # 解析为JSON格式
            if st.button("解析普通变量数据", key="parse_normal_vars"):
                try:
                    # 调用excel_to_json函数，使用空tables_data_mapping和指定sheetname
                    all_data = excel_to_json(uploaded_excel, "A1", {}, sheetname=selected_sheet)
                    
                    # 显示解析后的数据结构
                    with st.sidebar.expander("查看解析后的数据结构", expanded=False):
                        st.success(f"成功解析 {len(all_data)} 条普通变量数据")
                        st.json(all_data)
                    return all_data
                except Exception as e:
                    st.error(f"数据解析失败：{e}")
                    import traceback
                    st.code(traceback.format_exc())
                    
        
        except Exception as e:
            st.error(f"Excel读取失败：{e}")
            import traceback
            st.code(traceback.format_exc())
            
            # 清理临时文件
            temp_excel_path = os.path.join(tempfile.gettempdir(), uploaded_excel.name)
            if os.path.exists(temp_excel_path):
                os.remove(temp_excel_path)

# 有表格变量一个sheet一个word
def with_table_vars():
    """
    处理带有表格变量的情况：用户上传Excel后，默认处理所有sheet，最终解析为JSON格式
    """
    # 上传Excel文件
    uploaded_excel = st.file_uploader("上传Excel文件（包含普通变量和表格变量）", type=["xlsx", "xls"], key="with_table_vars_excel")
    
    if uploaded_excel:
        try:
            # 保存上传的Excel文件到临时目录
            temp_excel_path = os.path.join(tempfile.gettempdir(), uploaded_excel.name)
            with open(temp_excel_path, "wb") as f:
                f.write(uploaded_excel.getvalue())
            
            # 普通变量区域设置
            st.subheader("普通变量区域")
            text_data_start = st.text_input("普通变量左上角单元格（例如：A1）", value="A1", key="text_data_start")
            
            # 表格变量区域设置
            st.subheader("表格变量区域")
            
            # 动态添加表格配置
            num_tables = st.number_input("表格数量", min_value=1, value=1, step=1, key="num_tables")
            
            tables_data_mapping = {}
            for i in range(num_tables):
                col1, col2 = st.columns(2)
                with col1:
                    table_name = st.text_input(f"表格{i+1}名称（例如：家庭成员）", value=f"{i+1}", key=f"table_name_{i}")
                with col2:
                    table_start = st.text_input(f"表格{i+1}左上角单元格（例如：A10）", value=f"A{10+i*10}", key=f"table_start_{i}")
                
                if table_name and table_start:
                    tables_data_mapping[table_name] = table_start
            
            # 解析按钮
            if st.button("解析所有工作表数据", key="parse_with_table_vars"):
                try:
                    # 调用本地的excel_to_json函数
                    all_data = excel_to_json(temp_excel_path, text_data_start, tables_data_mapping)
                    
                    # 显示解析后的数据结构
                    with st.sidebar.expander("查看解析后的数据结构", expanded=False):
                        st.success(f"成功解析 {len(all_data)} 个工作表的数据")
                        st.json(all_data)
                    
                    # 删除临时文件
                    os.remove(temp_excel_path)
                    
                    # 直接返回解析结果
                    return all_data
                    
                except Exception as e:
                    st.error(f"数据解析失败：{e}")
                    import traceback
                    st.code(traceback.format_exc())
                    
                    # 清理临时文件
                    if os.path.exists(temp_excel_path):
                        os.remove(temp_excel_path)
        
        except Exception as e:
            st.error(f"Excel读取失败：{e}")
            import traceback
            st.code(traceback.format_exc())

# json解析
def parse_json():
    # JSON输入框
    json_input = st.text_area("输入JSON数据（支持多个家庭，用数组包裹）", height=150, placeholder='例如：[{"text_data": {...}, "tables_data": {...}}]')
    if json_input:
        try:
            result = json.loads(json_input)
            # 确保是数组格式
            if not isinstance(result, list):
                result = [result]
            # 显示解析后的数据结构
            with st.sidebar.expander("查看解析后的数据结构", expanded=False):
                st.success(f"成功解析 {len(result)} 条数据")
                st.json(result)
            return result
        except json.JSONDecodeError as e:
            st.error(f"JSON解析失败：{e}")

# -------------------------- Streamlit应用 --------------------------
if __name__ == "__main__":
    # 设置页面标题和布局
    st.set_page_config(page_title="Word模板批量生成", layout="wide")
    
    # 初始化session_state
    if "all_data" not in st.session_state:
        st.session_state.all_data = []
    
    # 主标题（使用markdown格式）
    st.markdown("### Word模板批量生成")
    
    # 调用模板函数
    cebianlan()
    
    # Word模板上传（使用session_state保存模板信息）
    st.subheader("📄 1、Word模板上传")
    new_uploaded_template = st.file_uploader("请上传Word模板文件(.docx)", type="docx")

    # 在侧边栏显示模板信息
    with st.sidebar:
        display_template_info(new_uploaded_template)
    
    # 主内容区只显示数据输入部分
    st.markdown("### 📊 2、数据输入")
    # 数据输入方式选择
    input_tab1, input_tab2, input_tab3 = st.tabs(["只有普通变量","有表格变量一个sheet一个word","JSON格式"])

    # 只有普通变量（一个sheet表生成多个word）
    with input_tab1:
        result = only_normal_vars()
        if result:
            st.session_state.all_data = result
    # 有表格变量一个sheet一个word
    with input_tab2:
        result = with_table_vars()
        if result:
            st.session_state.all_data = result
    
    # JSON输入方式
    with input_tab3:
        result = parse_json()
        if result:
            st.session_state.all_data = result

    # 批量生成文档按钮
    st.write("---")
    st.subheader("🏭 3、生成文档")
    # 批量生成文档按钮
    if st.button("批量生成文档"):
        st.write("🚀 开始...")
        st.code(st.session_state.all_data)
        st.write(new_uploaded_template)
        
        if st.session_state.all_data and new_uploaded_template:
            st.write("🚀 开始生成文档...")
            
            # 创建生成的文档列表
            generated_docs = []
            
            # 创建一个固定高度的容器，用于显示滚动消息
            with st.container(height=200, border=True):
                for index, family in enumerate(st.session_state.all_data, start=1):
                    try:
                        doc_stream, file_name = generate_contract_doc(new_uploaded_template, family, doc_index=index)
                        generated_docs.append((doc_stream, file_name))
                        st.success(f"✅ {file_name} 生成成功")
                    except Exception as e:
                        st.error(f"❌ 生成 {index} 号文档失败：{e}")
                        import traceback
                        st.code(traceback.format_exc())
            
            # 提供下载链接
            if generated_docs:
                st.write("---")
                st.subheader("下载生成的文档")
                
                # 创建ZIP文件
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    for doc_stream, file_name in generated_docs:
                        # 将文档流写入ZIP文件
                        zip_file.writestr(file_name, doc_stream.getvalue())
                zip_buffer.seek(0)
                
                # 提供ZIP下载
                st.download_button(
                    label=f"📦 批量下载所有文档 ({len(generated_docs)}个)",
                    data=zip_buffer,
                    file_name="批量生成文档.zip",
                    mime="application/zip",
                    use_container_width=True
                )
                
                # 保留单个文件下载选项
                st.markdown("<div style='margin-top: 20px; font-size: 14px; color: #666;'>或单独下载：</div>", unsafe_allow_html=True)
                for doc_stream, file_name in generated_docs:
                    st.download_button(
                        label=f"下载 {file_name}",
                        data=doc_stream,
                        file_name=file_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        
    
    # 在侧边栏添加清理按钮（可选）
    with st.sidebar:
        st.write("---")
        if st.button("清理临时文件"):
            # 清理st.session_state.all_data
            st.session_state.all_data = []
            st.sidebar.success("临时数据已清理")