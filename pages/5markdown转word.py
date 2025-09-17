# md2docx_app.py
import streamlit as st
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown, re, io, base64, pathlib, os
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import requests

# ---------- 常量：Markdown标签级别映射 ----------
LV_MAP = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}
LV_CN = {1: "标题 1", 2: "标题 2", 3: "标题 3", 4: "标题 4", 5: "标题 5", 6: "标题 6", 0: "正文"}

@st.cache_data(show_spinner=False)
def md2docx_bytes(md_text: str, template_name="数字章节", custom_template=None) -> io.BytesIO:
    """把 Markdown 文本转成 DOCX 并返回内存文件"""
    # 获取当前脚本所在目录
    current_dir = os.path.dirname(os.path.abspath(__file__))
    # 构建static文件夹的绝对路径
    static_dir = os.path.join(current_dir, "..", "static")
    
    # 处理选择了自定义模板但未上传文件的情况
    if template_name == "自定义模板" and custom_template is None:
        raise ValueError("请上传自定义模板文件！")
    
    # 优先使用自定义上传的模板
    if custom_template is not None:
        try:
            doc = Document(custom_template)
        except Exception as e:
            # 加载失败时回退到默认模板（数字章节）
            template_path = os.path.join(static_dir, "数字章节.docx")
            doc = Document(template_path)
    else:
        # 如果没有自定义模板，则使用选择的预设模板
        if template_name == "汉字章节":
            template_path = os.path.join(static_dir, "汉字章节.docx")
        else:
            template_path = os.path.join(static_dir, "数字章节.docx")
        doc = Document(template_path)

    # 检查并确保必要的样式存在
    def ensure_style_exists(style_name):
        try:
            # 尝试访问样式
            doc.styles[style_name]
            return True
        except KeyError:
            return False

    # 解析成 HTML
    html = markdown.markdown(md_text, extensions=['extra', 'codehilite', 'toc'])
    soup = BeautifulSoup(html, 'html.parser')

    def add_para(text, style=None, bold=False, italic=False):
        # 如果指定了样式但该样式不存在，则不使用样式
        if style and not ensure_style_exists(style):
            style = None
        
        p = doc.add_paragraph(style=style)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        return p

    for tag in soup:
        if tag.name == 'h1':
            doc.add_heading(tag.get_text(), level=1)
        elif tag.name == 'h2':
            doc.add_heading(tag.get_text(), level=2)
        elif tag.name == 'h3':
            doc.add_heading(tag.get_text(), level=3)
        elif tag.name in ('h4', 'h5', 'h6'):
            doc.add_heading(tag.get_text(), level=4)
        elif tag.name == 'p':
            # 简单处理行内 ** ** * *
            txt = tag.get_text()
            if '**' not in txt and '*' not in txt:
                add_para(txt)
            else:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*|\*[^*]*\*)', txt)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        p.add_run(part[2:-2]).bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        p.add_run(part[1:-1]).italic = True
                    else:
                        p.add_run(part)
        elif tag.name == 'ul':
            for li in tag.find_all('li', recursive=False):
                # 检查样式是否存在，如果不存在则创建项目符号
                if ensure_style_exists('List Bullet'):
                    add_para(li.get_text(), style='List Bullet')
                else:
                    p = add_para('• ' + li.get_text())
        elif tag.name == 'ol':
            for i, li in enumerate(tag.find_all('li', recursive=False), 1):
                # 检查样式是否存在，如果不存在则创建数字列表
                if ensure_style_exists('List Number'):
                    add_para(li.get_text(), style='List Number')
                else:
                    p = add_para(f'{i}. ' + li.get_text())
        elif tag.name == 'blockquote':
            p = add_para(tag.get_text())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.left_indent = Cm(0.8)
            p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif tag.name == 'pre':
            # 代码块
            p = add_para(tag.get_text())
            p.runs[0].font.name = 'Consolas'
            from docx.oxml.shared import OxmlElement, qn
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F5F5F5')
            p.paragraph_format.element.get_or_add_pPr().append(shading_elm)
        elif tag.name == 'table':
            # 简单表格
            rows = tag.find_all('tr')
            if not rows:
                continue
            cols = len(rows[0].find_all('th') or rows[0].find_all('td'))
            table = doc.add_table(rows=1, cols=cols)
            # 检查表格样式是否存在
            if ensure_style_exists('Table Grid'):
                table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, th in enumerate(rows[0].find_all('th')):
                hdr_cells[i].text = th.get_text()
            for tr in rows[1:]:
                row_cells = table.add_row().cells
                for i, td in enumerate(tr.find_all('td')):
                    row_cells[i].text = td.get_text()
        elif tag.name == 'img':
            # 网络图片直接下载嵌入
            src = tag.get('src')
            if src and src.startswith('http'):
                try:
                    response = st.session_state.get('__img_resp', None)
                    if response is None:
                        response = requests.get(src, timeout=10)
                        st.session_state['__img_resp'] = response
                    img_io = io.BytesIO(response.content)
                    doc.add_picture(img_io, width=Cm(12))
                except Exception as e:
                    st.warning(f"图片插入失败：{e}")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 修改获取用户输入的函数，拆分为上传和编辑两个部分
def get_file_upload():
    """获取上传的Markdown文件"""
    uploaded = st.file_uploader("上传Markdown文件", type=['md'])
    file_name = "converted.docx"
    md_text = None
    
    if uploaded is not None:
        md_text = uploaded.read().decode('utf-8')
        file_name = uploaded.name.replace('.md', '.docx')
        # 将上传的内容设置到文本编辑框中
        st.session_state.md_text = md_text
        st.session_state.file_name = file_name
    
    return md_text, file_name

def get_text_input():
    """获取用户编辑的Markdown内容"""
    # 使用session_state保存输入内容，确保上传文件后能更新编辑框
    md_text = st.text_area("编辑Markdown内容", value=st.session_state.get('md_text', ''), height=300)
    
    # 更新session_state中的内容
    if md_text != st.session_state.get('md_text', ''):
        st.session_state.md_text = md_text
        # 如果没有指定文件名，生成一个默认的
        if 'file_name' not in st.session_state:
            st.session_state.file_name = "converted.docx"
    
    return md_text, st.session_state.get('file_name', 'converted.docx')

def show_sidebar():
    """展示侧边栏推荐网站"""
    with st.sidebar:
        st.caption("### 🌟 md转word网站")
        st.markdown("- [markdowntoword](https://markdowntoword.io/zh)")
        st.markdown("- [cloudconvert](https://cloudconvert.com/md-to-docx)")
        st.markdown("- [convertio](https://convertio.co/zh/md-docx/)")
        st.caption("### 🌟 word转md网站")
        st.markdown("- [word转md](https://www.word2md.net/zh)")


def get_template_option():
    """获取用户选择的模板选项"""
    # 获取当前脚本所在目录
    current_dir = os.path.dirname(os.path.abspath(__file__))
    # 构建static文件夹的绝对路径
    static_dir = os.path.join(current_dir, "..", "static")
    
    template_option = st.radio(
        "请选择要使用的模板样式：",
        ["数字章节", "汉字章节", "自定义模板"],
        index=0,
        horizontal=True
    )
    szzj_path = os.path.join(static_dir, "szzj.png")
    st.image(szzj_path, caption="数字章节模板", width=250)
    hzzj_path = os.path.join(static_dir, "hzzj.png")
    st.image(hzzj_path, caption="汉字章节模板", width=250)
        
    custom_template = None
    # 只有选择自定义模板选项时，才显示上传按钮
    if template_option == "自定义模板":
        custom_template = st.file_uploader(
            "上传 .docx 格式的自定义模板",
            type="docx",
            help="上传自定义的Word模板文档"
        )
    
    return template_option, custom_template


def show_real_time_preview(md_text):
    """展示Markdown内容实时预览"""
    if md_text:
        try:
            # 转换Markdown为HTML进行预览
            html = markdown.markdown(md_text, extensions=['extra', 'codehilite', 'toc'])
            st.components.v1.html(html, height=500, scrolling=True)
        except Exception as e:
            st.error(f"预览失败：{e}")
    else:
        st.info("请在上方编辑Markdown内容，这里将显示实时预览效果。")


def process_conversion(md_text, file_name, template_option, custom_template):
    """处理Markdown转换并提供下载"""
    if st.button("生成 Word 文件", use_container_width=True, type="primary") and md_text:
        with st.spinner(f"使用'{template_option}'模板生成Word文件中..."):
            try:
                docx_buffer = md2docx_bytes(md_text, template_option, custom_template)
                st.success("转换完成！")
                st.download_button(
                    label="⬇️ 下载 Word 文件",
                    data=docx_buffer,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"转换失败：{e}")


if __name__ == "__main__":
    # 初始化session_state
    if 'md_text' not in st.session_state:
        st.session_state.md_text = ''
    if 'file_name' not in st.session_state:
        st.session_state.file_name = 'converted.docx'
    
    # 设置页面配置为宽屏布局
    st.set_page_config(page_title="MD→DOCX 在线转换", layout="wide")
    st.markdown("### 📄 Markdown 转 Word 工具")
    # 显示侧边栏
    show_sidebar()
    
    # 创建左右两列布局
    left_col, right_col = st.columns(2)
    
    with left_col:
        # 左上方：上传Markdown文件区域
        uploaded_md_text, uploaded_file_name = get_file_upload()
        # 左下方：编辑Markdown内容区域
        md_text, file_name = get_text_input()
    
    with right_col:
        # 获取模板选项
        template_option, custom_template = get_template_option()
    
    # 处理转换和下载
    process_conversion(st.session_state.md_text, st.session_state.file_name, template_option, custom_template)
    # 实时预览区域
    show_real_time_preview(st.session_state.md_text)