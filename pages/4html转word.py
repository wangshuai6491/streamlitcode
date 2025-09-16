import streamlit as st
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
import pandas as pd
import io
import base64
import os
# ---------- 常量：HTML标签级别映射，样式 ----------
LV_MAP = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}
LV_CN = {1: "标题 1", 2: "标题 2", 3: "标题 3", 4: "标题 4", 5: "标题 5", 6: "标题 6", 0: "正文"}

# 默认内容标签列表
DEFAULT_CONTENT_TAGS = ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'span']

# ---------- 函数 ----------
def quick_tag_map(html: str, content_tags):
    """快速映射HTML标签及其示例文本"""
    soup = BeautifulSoup(html, "html.parser")
    tag_sample = {}
    # 只收集内容标签
    for tag in soup.find_all(content_tags):
        text = tag.get_text(" ", strip=True)
        if text and tag.name not in tag_sample:
            tag_sample[tag.name] = text[:60]
    # 创建表格数据，只包含标签和示例，级别信息在主流程中添加
    rows = [{"标签": f"<{k}>", "示例": v, "default_level": LV_MAP.get(k, 0)} for k, v in tag_sample.items()]
    return pd.DataFrame(rows)

# 存储原始标签信息的字典
tag_info_map = {}

def process_element(element, content_tags, blocks):
    """递归处理HTML元素，构建文档块"""
    # 处理文本节点
    if element.name is None:
        text = element.strip()
        if text:
            return text
        return ''
    # 跳过script和style标签
    elif element.name in ['script', 'style']:
        return ''
    # 处理内容标签
    elif element.name in content_tags:
        text = ''
        has_content_child = False
        
        # 检查是否有子内容标签
        for child in element.children:
            if child.name in content_tags:
                has_content_child = True
                # 递归处理子内容标签
                process_element(child, content_tags, blocks)
            else:
                # 处理非内容标签的子元素
                child_text = process_element(child, content_tags, blocks) or ''
                text += child_text
        
        # 如果没有子内容标签，才将当前内容标签添加到blocks
        if not has_content_child:
            full_text = text.strip()
            if full_text:
                block_id = len(blocks)
                blocks.append({
                    "text": full_text,
                    "level": LV_MAP.get(element.name, 0),
                    "tag": element.name  # 保存原始标签名称
                })
                # 存储标签信息映射，用于预览时精确匹配
                tag_info_map[block_id] = element.name
        
        return ''
    # 处理其他标签（如div等容器标签）
    else:
        text = ''
        for child in element.children:
            child_text = process_element(child, content_tags, blocks) or ''
            text += child_text
        return text

def parse_html_to_blocks(html_raw, content_tags):
    """将HTML解析为文档块列表"""
    soup = BeautifulSoup(html_raw, "html.parser")
    blocks = []
    global tag_info_map
    tag_info_map = {}
    
    # 递归处理HTML元素
    process_element(soup, content_tags, blocks)
    
    # 如果没有找到任何内容标签，尝试直接从body中提取文本
    if not blocks:
        body = soup.find('body')
        if body:
            text = body.get_text(" ", strip=True)
            if text:
                blocks.append({"text": text, "level": 0})
    
    return blocks

def build_docx(blocks, template_name="数字章节", custom_template=None):
    """根据文档块使用自定义模板构建DOCX文档"""
    # 获取当前脚本所在目录
    current_dir = os.path.dirname(os.path.abspath(__file__))
    # 构建static文件夹的绝对路径
    static_dir = os.path.join(current_dir, "..", "static")
    
    # 处理选择了自定义模板但未上传文件的情况
    if template_name == "自定义模板" and custom_template is None:
        st.error("请上传自定义模板文件！")
        return None
    
    # 优先使用自定义上传的模板
    if custom_template is not None:
        try:
            doc = Document(custom_template)
        except Exception as e:
            st.error(f"自定义模板加载失败：{str(e)}")
            # 加载失败时回退到默认模板（数字章节）
            template_path = os.path.join(static_dir, "数字章节.docx")
            doc = Document(template_path)
    else:
        # 如果没有自定义模板，则使用选择的预设模板
        if template_name == "汉字章节":
            template_path = os.path.join(static_dir, "汉字章节.docx")
        else:
            template_path = os.path.join(static_dir, "数字章节.docx")
        doc = Document(template_path)
    
    for blk in blocks:
        level = blk["level"]
        para = doc.add_paragraph()
        run = para.add_run(blk["text"])
        
        # 使用模板中的样式
        if level >= 1:
            para.style = f"Heading {level}"
    
    return doc

def create_download_link(doc):
    """创建DOCX文件的下载链接"""
    if doc is None:
        return None
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    b64 = base64.b64encode(buffer.read()).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="export.docx">⏬ 下载 DOCX</a>'
    return href


def get_main_input():
    """获取主页面用户输入"""
    # 创建两列布局，左侧占更大空间
    col1, col2 = st.columns([3, 1])  # 左侧宽度是右侧的3倍
    
    with col1:
        st.markdown("### 0.📝 HTML 输入区域")
        html_raw = st.text_area("粘贴 HTML", height=150, help="在此粘贴您想要转换的HTML代码")
    
    with col2:
        # 添加自定义内容标签配置
        st.markdown("### 🔧 自定义内容标签")
        default_tags_str = ", ".join(DEFAULT_CONTENT_TAGS)
        custom_tags = st.text_input("作为独立段落的标签（逗号分隔）", value=default_tags_str, 
                                    help="例如: h1, h2, p, li, span, div")
        content_tags = [tag.strip() for tag in custom_tags.split(",") if tag.strip()]
    
    # 解析按钮放在下方
    parsed = False
    if st.button("🔍 开始解析", type="primary"):
        if html_raw:
            parsed = True
        else:
            st.warning("请先粘贴HTML代码")
    
    return html_raw, content_tags, parsed

def show_recommended_sites():
    """显示推荐网站侧边栏"""
    with st.sidebar:
        st.caption("### 🌟 推荐网站")
        st.markdown("- [kalvinbg](https://tools.kalvinbg.cn/doc/html2word)")
        st.markdown("- [freeconvert](https://www.freeconvert.com/zh/html-to-word)")
        st.markdown("- [cdkm](https://cdkm.com/cn/html-to-doc)")
        
        # 分隔线
        st.markdown("---")
        st.markdown("💡 希望这些网站对您有帮助！")

# ---------- 主流程 ----------
if __name__ == "__main__":
    st.set_page_config(page_title="HTML → DOCX 精调导出", layout="wide")
    # 显示推荐网站侧边栏
    show_recommended_sites()
    
    # 初始化会话状态
    if "parsed_data" not in st.session_state:
        st.session_state["parsed_data"] = None
    if "tag_lv" not in st.session_state:
        st.session_state["tag_lv"] = None
    
    # 获取主页面输入
    html_raw, content_tags, parsed = get_main_input()
    
    # 如果点击了解析按钮
    if parsed:
        # 解析HTML结构，生成标签映射表
        df_map = quick_tag_map(html_raw, content_tags)
        
        # 解析HTML为文档块
        blocks = parse_html_to_blocks(html_raw, content_tags)
        
        # 保存解析结果到会话状态
        st.session_state["parsed_data"] = {
            "df_map": df_map,
            "blocks": blocks
        }
        
        # 初始化标签级别映射
        df = df_map.copy()
        df.insert(1, "级别", df["default_level"].map(LV_CN))
        rev_cn = {v: k for k, v in LV_CN.items()}
        st.session_state["tag_lv"] = {row["标签"]: rev_cn[row["级别"]] for _, row in df.iterrows()}

    # 如果已经解析过数据，显示功能界面
    if st.session_state["parsed_data"]:
        # 获取解析数据
        df = st.session_state["parsed_data"]["df_map"].copy()
        blocks = st.session_state["parsed_data"]["blocks"].copy()
        
        # 1. 调整HTML标签对应的文档级别
        st.markdown("### 1.调整HTML标签对应的文档级别")
        # 添加级别列，并设置为可编辑
        df.insert(1, "级别", df["default_level"].map(LV_CN))
        df = df.drop(columns=["default_level"])
        
        edited_df = st.data_editor(
            df,
            use_container_width=True,
            column_config={"级别": st.column_config.SelectboxColumn(options=list(LV_CN.values()))},
            disabled=["标签", "示例"],
            key="editor"
        )
        
        # 保存映射结果到会话状态
        rev_cn = {v: k for k, v in LV_CN.items()}
        st.session_state["tag_lv"] = {row["标签"]: rev_cn[row["级别"]] for _, row in edited_df.iterrows()}
        
        # 2. 样式选择
        st.markdown("### 2.选择模板样式")
        # 获取当前脚本所在目录
        current_dir = os.path.dirname(os.path.abspath(__file__))
        # 构建static文件夹的绝对路径
        static_dir = os.path.join(current_dir, "..", "static")
        
        col1, col2 = st.columns(2)
        
        with col1:
            szzj_path = os.path.join(static_dir, "szzj.png")
            st.image(szzj_path, caption="数字章节模板", width=300)
        with col2:
            hzzj_path = os.path.join(static_dir, "hzzj.png")
            st.image(hzzj_path, caption="汉字章节模板", width=300)
            
        template_option = st.radio(
            "请选择要使用的模板样式：",
            ["数字章节", "汉字章节", "自定义模板"],
            index=0,
            horizontal=True
        )
        
        custom_template = None
        # 只有选择自定义模板选项时，才显示上传按钮
        if template_option == "自定义模板":
            custom_template = st.file_uploader(
                "上传 .docx 格式的自定义模板",
                type="docx",
                help="上传自定义的Word模板文档"
            )
        
        # 3. 下载功能
        rev_cn = {v: k for k, v in LV_CN.items()}
        if st.button("📄 生成 DOCX 文档", type="primary"):
            with st.spinner(f"使用'{template_option}'模板生成DOCX文档中..."):
                # 创建新的blocks副本以避免修改原始数据
                updated_blocks = []
                for blk in blocks:
                    updated_blk = blk.copy()
                    # 使用原始标签信息进行精确匹配
                    original_tag = blk.get("tag")
                    if original_tag:
                        # 查找原始标签对应的映射级别
                        for tag_name, lvl in st.session_state["tag_lv"].items():
                            if tag_name.strip("<>") == original_tag:
                                updated_blk["level"] = lvl
                                break
                    updated_blocks.append(updated_blk)
                
                # 构建并提供下载链接
                doc = build_docx(updated_blocks, template_option, custom_template)
                if doc is not None:
                    download_link = create_download_link(doc)
                    if download_link is not None:
                        st.markdown(download_link, unsafe_allow_html=True)
            