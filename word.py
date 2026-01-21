from docx import Document
import os
import re


def analyze_template(template_path):
    """
    分析Word模板，提取所有占位符信息
    
    :param template_path: Word模板文件路径(.docx格式)
    :return: 包含普通变量和表格信息的字典
    """
    # 加载Word模板
    doc = Document(template_path)
    
    # 正则表达式匹配占位符 {xxx}
    placeholder_pattern = re.compile(r'{([^{}]*)}')
    all_placeholders = set()
    
    # 遍历所有段落，提取占位符
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            matches = placeholder_pattern.findall(run.text)
            for match in matches:
                all_placeholders.add(f"{{{match}}}")
    
    # 遍历所有表格，提取占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = placeholder_pattern.findall(cell.text)
                for match in matches:
                    all_placeholders.add(f"{{{match}}}")
    
    # 分类占位符：普通变量和表格变量
    normal_placeholders = []
    table_info = {}
    
    for ph in all_placeholders:
        ph_content = ph[1:-1]  # 去掉 {} 括号
        if ph_content.startswith("table."):
            # 分割表格占位符：table.表格名.列名
            parts = ph_content.split(".")
            if len(parts) == 3:
                table_name = parts[1]
                column_name = parts[2]
                
                # 添加到表格信息字典
                if table_name not in table_info:
                    table_info[table_name] = {
                        "columns": set(),
                        "placeholders": []
                    }
                table_info[table_name]["columns"].add(column_name)
                table_info[table_name]["placeholders"].append(ph)
        else:
            normal_placeholders.append(ph)
    
    # 转换columns集合为列表，方便使用
    for table_name in table_info:
        table_info[table_name]["columns"] = sorted(list(table_info[table_name]["columns"]))
    
    return {
        "normal_placeholders": sorted(normal_placeholders),
        "table_info": table_info
    }


def generate_contract_doc(template_path, output_dir, data, table_index=None, doc_index=1):
    """
    基于Word模板生成文档
    
    :param template_path: Word模板文件路径(.docx格式)
    :param output_dir: 生成文档的保存目录
    :param data: 文档数据，字典格式，包含text_data和tables_data
    :param table_index: 可选参数，指定表格索引（0开始），跳过自动识别
    :param doc_index: 文档序号，用于生成文件名
    
    """
    # 1. 创建保存目录（如果不存在）
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 2. 加载Word模板
    doc = Document(template_path)

    # 3. 分析模板，获取占位符信息
    template_analysis = analyze_template(template_path)
    
    # 打印普通变量占位符
    print("📋 普通变量占位符：")
    for ph in template_analysis["normal_placeholders"]:
        print(f"   - {ph}")
    
    # 打印表格信息
    print("📋 表格信息：")
    for table_name, info in template_analysis["table_info"].items():
        print(f"   📊 表格名称：{table_name}")
        print(f"      列名：{', '.join(info['columns'])}")
        print(f"      占位符：{', '.join(info['placeholders'])}")

    # ========== 核心功能1：替换所有【普通文本占位符】 {xxx} ==========
    # 遍历文档中所有段落，替换文本占位符
    for paragraph in doc.paragraphs:
        # 遍历段落中的所有文本块，避免替换格式丢失
        for run in paragraph.runs:
            for placeholder, value in data["text_data"].items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

    # ========== 核心功能2：动态填充【表格】并插入多行 ==========
    # 遍历所有表格，处理包含table.xxx.xxx格式占位符的表格
    for table in doc.tables:
        # 检查表格是否包含模板行（包含table.xxx.xxx占位符）
        template_row = None
        table_name_in_row = None
        
        for row_idx, row in enumerate(table.rows):
            row_text = " ".join([cell.text for cell in row.cells])
            if "{table." in row_text:
                # 提取表格名称
                matches = re.findall(r'{table\.([^\.}]+)\.', row_text)
                if matches:
                    table_name_in_row = matches[0]
                    template_row = row
                    break
        
        if template_row and table_name_in_row:
            # 检查是否有对应的数据
            if table_name_in_row not in data.get("tables_data", {}):
                print(f"⚠️  表格 {table_name_in_row} 没有对应的数据，跳过该表格")
                continue
            
            # 获取表格数据
            table_data = data["tables_data"][table_name_in_row]
            
            # 获取模板行中各占位符的位置映射
            placeholder_positions = {}
            template_cells = template_row.cells
            for i, cell in enumerate(template_cells):
                cell_text = cell.text.strip()
                if cell_text.startswith("{table.") and cell_text.endswith("}"):
                    placeholder_positions[cell_text] = i
            
            # 遍历表格数据，逐行插入
            for index, row_data in enumerate(table_data):
                if index == 0:
                    # 第一行数据：直接使用模板行
                    row = template_row.cells
                    # 保存第一行的单元格作为格式模板
                    format_template_cells = template_row.cells
                else:
                    # 第2-N行数据：在模板行下方新增一行
                    new_row = table.add_row()
                    row = new_row.cells
                    
                    # 复制第一行的格式到新行
                    # 复制行属性
                    new_row.height = template_row.height
                    
                    # 复制每个单元格的格式
                    for i, (new_cell, template_cell) in enumerate(zip(row, format_template_cells)):
                        # 复制单元格宽度
                        new_cell.width = template_cell.width
                        
                        # 复制段落格式
                        for new_paragraph, template_paragraph in zip(new_cell.paragraphs, template_cell.paragraphs):
                            # 复制段落对齐方式
                            new_paragraph.alignment = template_paragraph.alignment
                            
                            # 复制段落样式
                            if template_paragraph.style:
                                new_paragraph.style = template_paragraph.style
                            
                            # 复制段落中的运行元素格式
                            for new_run, template_run in zip(new_paragraph.runs, template_paragraph.runs):
                                # 复制字体格式
                                new_run.font.name = template_run.font.name
                                new_run.font.size = template_run.font.size
                                new_run.font.bold = template_run.font.bold
                                new_run.font.italic = template_run.font.italic
                                new_run.font.underline = template_run.font.underline
                                new_run.font.color.rgb = template_run.font.color.rgb
                                
                                # 复制其他字体属性
                                new_run.font.superscript = template_run.font.superscript
                                new_run.font.subscript = template_run.font.subscript
                                new_run.font.shadow = template_run.font.shadow
                                new_run.font.strike = template_run.font.strike
                                new_run.font.double_strike = template_run.font.double_strike
                
                # 给当前行的每个单元格赋值
                for placeholder, pos in placeholder_positions.items():
                    ph_content = placeholder[1:-1]  # 去掉 {} 括号
                    # 分割占位符：table.表格名.列名
                    parts = ph_content.split(".")
                    if len(parts) == 3:
                        column_name = parts[2]
                        
                        if column_name == "序号":
                            row[pos].text = str(index + 1)  # 序号从1开始
                        elif column_name in row_data:
                            row[pos].text = str(row_data[column_name])
                        else:
                            # 如果数据中没有该列，留空或使用默认值
                            row[pos].text = ""

    # 4. 保存生成的文档
    # 提取提供的数据中有没有文件名
    file_name = data.get("text_data", {}).get("{文件名}", "")
    # 如果文件名为空或只有空格，则按顺序号命名
    if not file_name.strip():
        file_name = f"{doc_index}.docx"
    
    save_path = os.path.join(output_dir, file_name)
    doc.save(save_path)
    print(f"✅ 文档生成成功：{save_path}")


if __name__ == "__main__":
    # -------------------------- 配置参数（修改这里即可） --------------------------
    WORD_TEMPLATE_PATH = "static/家庭承包制模板.docx"  # 你的Word模板路径
    OUTPUT_DIR = "结果文档"              # 生成文档的保存目录

    # -------------------------- 测试数据（多户家庭，按需修改） --------------------------
    # 每户数据包含：text_data(文本占位符) + tables_data(所有表格数据)
    all_family_data = [
  {
    "text_data": {
      "{家庭编号}": "家庭1",
      "{户主姓名}": "张三",
      "{家庭住址}": "北京市朝阳区建国路88号",
      "{承包证号}": "京农承2026001",
      "{承包地块}": "村东头良田3块、村西坡地1块",
      "{承包面积}": "5.2",
      "{承包起始日期}": "2026-01-01",
      "{承包终止日期}": "2056-01-01",
      "{村委会名称}": "北京市朝阳区建国村委会",
      "{签署日期}": "2026-01-19",
      "{文件名}": ""
    },
    "tables_data": {
      "家庭成员": [
        {
          "家庭编号": "家庭1",
          "姓名": "张三",
          "性别": "男",
          "年龄": "45",
          "关系": "户主",
          "身份证": "110105198101011234",
          "是否承包": "是"
        },
        {
          "家庭编号": "家庭1",
          "姓名": "李梅",
          "性别": "女",
          "年龄": "42",
          "关系": "配偶",
          "身份证": "110105198402025678",
          "是否承包": "是"
        },
        {
          "家庭编号": "家庭1",
          "姓名": "张小明",
          "性别": "男",
          "年龄": "18",
          "关系": "长子",
          "身份证": "110105200803039012",
          "是否承包": "是"
        }
      ],
      "测试表": [
        {
          "家庭编号": "家庭1",
          "测试表A": "A1",
          "测试表B": "B1",
          "测试表C": "C1"
        },
        {
          "家庭编号": "家庭1",
          "测试表A": "A2",
          "测试表B": "B2",
          "测试表C": "C2"
        },
        {
          "家庭编号": "家庭1",
          "测试表A": "A3",
          "测试表B": "B3",
          "测试表C": "C3"
        }
      ]
    }
  },
  {
    "text_data": {
      "{家庭编号}": "家庭2",
      "{户主姓名}": "王老太",
      "{家庭住址}": "上海市浦东新区张江路66号",
      "{承包证号}": "沪农承2026002",
      "{承包地块}": "村南自留地1块",
      "{承包面积}": "1.0",
      "{承包起始日期}": "2026-01-01",
      "{承包终止日期}": "2056-01-01",
      "{村委会名称}": "上海市浦东新区张江村委会",
      "{签署日期}": "2026-01-19",
      "{文件名}": "王老太的.docx"
    },
    "tables_data": {
      "家庭成员": [
        {
          "家庭编号": "家庭2",
          "姓名": "王老太",
          "性别": "女",
          "年龄": "78",
          "关系": "户主",
          "身份证": "310115194805056789",
          "是否承包": "是"
        }
      ],
      "测试表": [
        {
          "家庭编号": "家庭2",
          "测试表A": 1,
          "测试表B": 3,
          "测试表C": 5
        },
        {
          "家庭编号": "家庭2",
          "测试表A": 2,
          "测试表B": 4,
          "测试表C": 6
        }
      ]
    }
  },
  {
    "text_data": {
      "{家庭编号}": "家庭3",
      "{户主姓名}": "刘建国",
      "{家庭住址}": "广州市天河区天河路385号",
      "{承包证号}": "粤农承2026003",
      "{承包地块}": "村西水田5块、村北旱地2块",
      "{承包面积}": "8.7",
      "{承包起始日期}": "2026-01-01",
      "{承包终止日期}": "2056-01-01",
      "{村委会名称}": "广州市天河区天河村委会",
      "{签署日期}": "2026-01-19",
      "{文件名}": "刘建国的.docx"
    },
    "tables_data": {
      "家庭成员": [
        {
          "家庭编号": "家庭3",
          "姓名": "刘建国",
          "性别": "男",
          "年龄": "48",
          "关系": "户主",
          "身份证": "440106197808081234",
          "是否承包": "是"
        },
        {
          "家庭编号": "家庭3",
          "姓名": "陈秀兰",
          "性别": "女",
          "年龄": "45",
          "关系": "配偶",
          "身份证": "440106198109095678",
          "是否承包": "是"
        },
        {
          "家庭编号": "家庭3",
          "姓名": "刘志强",
          "性别": "男",
          "年龄": "20",
          "关系": "长子",
          "身份证": "440106200603039012",
          "是否承包": "是"
        },
        {
          "家庭编号": "家庭3",
          "姓名": "刘美丽",
          "性别": "女",
          "年龄": "16",
          "关系": "长女",
          "身份证": "440106201005053456",
          "是否承包": "否"
        },
        {
          "家庭编号": "家庭3",
          "姓名": "刘老夫",
          "性别": "男",
          "年龄": "72",
          "关系": "父亲",
          "身份证": "440106195402027890",
          "是否承包": "否"
        }
      ],
      "测试表": [
        {
          "家庭编号": "家庭3",
          "测试表A": "Q",
          "测试表B": "W",
          "测试表C": "E"
        }
      ]
    }
  },
  {
    "text_data": {
      "{家庭编号}": "家庭4",
      "{户主姓名}": "ASDA",
      "{家庭住址}": "广州市天河区天河路385号",
      "{承包证号}": "粤农承2026003",
      "{承包地块}": "村西水田5块、村北旱地2块",
      "{承包面积}": "8.7",
      "{承包起始日期}": "2026-01-01",
      "{承包终止日期}": "2056-01-01",
      "{村委会名称}": "广州市天河区天河村委会",
      "{签署日期}": "2026-01-19",
      "{文件名}": "a的.docx"
    },
    "tables_data": {
      "家庭成员": [
        {
          "家庭编号": "家庭4",
          "姓名": "FFG",
          "性别": "男",
          "年龄": 12,
          "关系": "户主",
          "身份证": "440106197808081234",
          "是否承包": "是"
        },
        {
          "家庭编号": "家庭4",
          "姓名": "WDF",
          "性别": "女",
          "年龄": 45,
          "关系": "配偶",
          "身份证": "440106198109095678",
          "是否承包": "是"
        }
      ],
      "测试表": [
        {
          "家庭编号": "家庭4",
          "测试表A": "Z",
          "测试表B": "R",
          "测试表C": "T"
        },
        {
          "家庭编号": "家庭4",
          "测试表A": "X",
          "测试表B": "F",
          "测试表C": "Y"
        },
        {
          "家庭编号": "家庭4",
          "测试表A": "C",
          "测试表B": "V",
          "测试表C": "U"
        }
      ]
    }
  }
]

    # -------------------------- 执行批量生成 --------------------------
    print("🚀 开始批量生成家庭承包经营权证...")
    for index, family in enumerate(all_family_data, start=1):
        generate_contract_doc(WORD_TEMPLATE_PATH, OUTPUT_DIR, family, doc_index=index)
    print("🎉 所有文档生成完成！")